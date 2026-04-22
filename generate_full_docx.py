import os
import glob
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Image directories ──
admin_img_dir = r"C:\Users\DAR-F\.gemini\antigravity\brain\b8436daf-f3e5-4bdd-89b5-22717f5c8b42"
student_img_dir = r"d:\ERU\lev4 sem1\graduiton project\edusphere fronend\rawda_front_v2\edusphere-project\r photo"
output_path = r"d:\ERU\lev4 sem1\graduiton project\edusphere fronend\rawda_front_v2\edusphere-project\admin panel pages docs.docx"

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

fig_counter = [1]

def find_img(pattern, directory=None):
    search_dir = directory or admin_img_dir
    files = glob.glob(os.path.join(search_dir, pattern))
    if not files:
        return None
    files.sort(key=os.path.getmtime, reverse=True)
    return files[0]

def add_page(pattern, title, description, fig_prefix="4", directory=None):
    img_path = find_img(pattern, directory)
    if not img_path:
        print(f"  WARNING: No image found for pattern '{pattern}' - skipping {title}")
        return
    doc.add_heading(title, level=3)
    try:
        doc.add_picture(img_path, width=Inches(6.0))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Image error: {e}]")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Figure {fig_prefix}.{fig_counter[0]} \u2013 {title}")
    run.italic = True
    run.font.size = Pt(10)
    fig_counter[0] += 1
    p_desc = doc.add_paragraph(description)
    p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()

def add_student_page(filename, title, description):
    """Shortcut for student portal images using exact filenames."""
    add_page(filename, title, description, directory=student_img_dir)


# ═══════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EduSphere University Portal")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Comprehensive System Documentation")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("All Roles: Student \u2022 Administrator \u2022 Instructor \u2022 Student Affairs")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
doc.add_page_break()

# ═══════════════════════════════════════════════
#  TABLE OF CONTENTS
# ═══════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Authentication",
    "2. Student Portal",
    "   2.1 Student Dashboard",
    "   2.2 Notifications Panel",
    "   2.3 Smart Attendance (QR)",
    "   2.4 Student Services",
    "   2.5 Curriculum Management",
    "   2.6 My Grades",
    "   2.7 Records",
    "   2.8 AI Assistant",
    "   2.9 Account Settings",
    "   2.10 Get Support & Sign Out",
    "   2.11 Dark Mode",
    "3. Administrator Role",
    "   3.1 Admin Dashboard",
    "   3.2 Student Management",
    "   3.3 Faculty Management",
    "   3.4 Course Management",
    "   3.5 Registration & Operations",
    "   3.6 Student Affairs Module",
    "   3.7 Scheduling",
    "   3.8 Academic Performance",
    "   3.9 Reports & Analytics",
    "   3.10 System Administration",
    "4. Instructor Role",
    "   4.1 Instructor Dashboard",
    "   4.2 Profile",
    "   4.3 My Courses",
    "   4.4 Grade Management",
    "   4.5 Reports",
    "5. Student Affairs Staff Role",
    "   5.1 Student Affairs Dashboard",
    "   5.2 Student Management",
    "   5.3 Student Affairs Operations",
    "   5.4 Registration Management",
    "   5.5 Reports",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
doc.add_page_break()

# ═══════════════════════════════════════════════
#  1. AUTHENTICATION
# ═══════════════════════════════════════════════
doc.add_heading("1. Authentication", level=1)
doc.add_paragraph(
    "The EduSphere University Portal requires authentication for all users. "
    "The system supports multiple distinct roles: Student, Administrator, Instructor (Faculty Member), "
    "and Student Affairs Staff. Each role has access to specific modules and features "
    "tailored to their responsibilities within the institution. The authentication module "
    "includes login, password recovery, and password reset functionalities to ensure "
    "secure and convenient access for all users."
)

add_student_page("WhatsApp Image 2026-04-02 at 6.50.13 PM.jpeg", "Login Page",
    "This is the Login page of the EduSphere University Portal. It features a modern split-screen design with "
    "institutional branding and a university building illustration on the left side, along with the EduSphere "
    "logo and 'University Portal' subtitle. The motivational tagline 'Empowering minds, shaping futures' is "
    "displayed at the bottom with a carousel indicator. The right panel presents a clean authentication form "
    "titled 'Welcome at EduSphere' with two input fields: 'Email' (with placeholder 'Enter here') and "
    "'Password' (with placeholder 'Enter password here' and a visibility toggle icon). A 'Remember me' "
    "checkbox allows users to persist their login session. The prominent red 'Sign In' button submits the "
    "credentials, and a 'Forget your password?' link below provides access to the password recovery flow. "
    "The interface supports role-based access control, automatically redirecting users to their appropriate "
    "dashboard upon successful authentication.")

add_student_page("WhatsApp Image 2026-04-02 at 6.50.13 PM (1).jpeg", "Forgot Password Page",
    "This is the Forgot Password page of the EduSphere University Portal. It maintains the consistent "
    "split-screen layout with institutional branding and the university illustration on the left side. "
    "The right panel contains a form titled 'Forgot Password' with the subtitle 'Enter your email to "
    "receive a reset code'. A single 'Email Address' input field with placeholder 'Enter your email' "
    "collects the user's registered email. The prominent red 'Send Reset Code' button dispatches a "
    "6-digit verification code to the provided email address. A '\u2190 Back to Login' link below the "
    "button allows users to return to the main login page.")

add_student_page("WhatsApp Image 2026-04-02 at 6.50.13 PM (2).jpeg", "Reset Password Page",
    "This is the Reset Password page, accessed after requesting a password reset code. It maintains the "
    "consistent split-screen layout. The right panel is titled 'Reset Password' with the subtitle 'Enter "
    "the code and your new password'. The form contains three input fields: 'Verification Code' (placeholder "
    "'Enter 6-digit code'), 'New Password' (placeholder 'Enter new password'), and 'Confirm Password' "
    "(placeholder 'Confirm new password'). The red 'Reset Password' button validates that both password "
    "fields match and the verification code is correct before updating the password. A '\u2190 Back to Login' "
    "link provides navigation back to the authentication page.")

doc.add_page_break()

# ═══════════════════════════════════════════════
#  2. STUDENT PORTAL
# ═══════════════════════════════════════════════
doc.add_heading("2. Student Portal", level=1)
doc.add_paragraph(
    "The Student Portal provides a personalized academic experience for enrolled students. "
    "It serves as the primary interface through which students interact with the university's "
    "digital ecosystem. Key features include a comprehensive dashboard with academic metrics, "
    "smart QR-based attendance tracking, access to student services such as medical excuses "
    "and official requests, curriculum management, academic records, an AI-powered assistant, "
    "and a fully customizable settings panel. The portal features a persistent left sidebar "
    "for navigation and a top bar with search, dark mode toggle, settings, notifications, and profile access."
)

# 2.1 Student Dashboard
doc.add_heading("2.1 Student Dashboard", level=2)

add_student_page("WhatsApp Image 2026-04-02 at 6.50.13 PM (3).jpeg", "Student Dashboard",
    "This is the main Student Dashboard, the landing page for authenticated students. The interface features "
    "a welcoming header 'Welcome back, Rawda \U0001f44b' with subtitle 'Here\u2019s what\u2019s happening with your "
    "studies today.' Quick action buttons for 'Timetable' and '+ Add Course' are in the top-right. Four KPI "
    "cards display: Total Courses (6, 'This Semester'), Average GPA (3.8 with green '+0.2 pts'), Classes "
    "Today (4, 'On track'), and Pending Tasks (8, '2 overdue' in red). The 'Today\u2019s Classes' section "
    "lists scheduled sessions with course names, instructors, rooms, times, and status badges \u2014 'Live Now' "
    "(red), 'Upcoming' (amber), 'Scheduled' (gray). A 'Recent Grades' panel shows CS431 Project 3 at 95%, "
    "MATH301 Midterm at 38%, ENG101 Essay at 92% with color-coded progress bars. The sidebar includes "
    "Dashboard, Attendance QR (badge 3), Student Services, Curriculum Management, Records, AI Assistant, "
    "Settings, Sign Out, and a 'Need Help? Get Support' section.")

doc.add_page_break()

# 2.2 Notifications Panel
doc.add_heading("2.2 Notifications Panel", level=2)

add_student_page("WhatsApp Image 2026-04-02 at 6.50.17 PM (2).jpeg", "Notifications Panel",
    "This screenshot shows the Student Dashboard with the Notifications side panel open, triggered by "
    "clicking the bell icon in the top navigation bar. The panel slides in from the right and displays "
    "real-time notifications organized chronologically. Visible notifications include: 'Class Starting Soon' "
    "\u2014 Advanced Web Development starts in 15 min (10 min ago), 'Grade Posted' \u2014 MATH301 Midterm "
    "Exam: 38% (3 hours ago), and 'Attendance Warning' \u2014 You missed 2 sessions in ENG101 (Yesterday). "
    "Each notification features a relevant icon, a bold title, descriptive text, and a timestamp. The panel "
    "includes a close button (\u00d7) in the top-right corner. This notification system keeps students informed "
    "about upcoming classes, newly posted grades, and attendance warnings in real-time.")

doc.add_page_break()

# 2.3 Smart Attendance
doc.add_heading("2.3 Smart Attendance (QR)", level=2)

add_student_page("WhatsApp Image 2026-04-02 at 6.50.13 PM (4).jpeg", "Smart Attendance \u2013 QR Scanner",
    "This is the Smart Attendance page, accessible from the 'Attendance QR' sidebar item. The header shows "
    "'Smart Attendance' with a back arrow and instruction 'Scan the QR code displayed in your lecture hall'. "
    "A real-time clock (08:16:31 PM) appears in the top-right. The main area features a camera preview with "
    "a 'Camera Ready' status and instruction 'Point at the QR code on the screen'. A full-width red 'Scan QR' "
    "button with camera icon initiates scanning. The system validates QR codes against the current session "
    "schedule, ensuring attendance can only be recorded during the correct time window, replacing traditional "
    "paper-based attendance with a tamper-proof digital solution.")

doc.add_page_break()

# 2.4 Student Services
doc.add_heading("2.4 Student Services", level=2)

add_student_page("WhatsApp Image 2026-04-02 at 6.50.13 PM (5).jpeg", "Student Services Hub",
    "This is the Student Services hub page, providing centralized access to all university services 'in one "
    "place' (highlighted in red). Four service cards are arranged in a 2\u00d72 grid: 'Medical Excuses' (green "
    "grid icon) for submitting medical documentation, 'Complaints' (red speech bubble icon) for filing "
    "complaints to administration, 'Academic Warnings' (yellow triangle icon) for viewing and appealing "
    "warnings, and 'Official Requests' (purple question mark icon) for requesting transcripts and certificates. "
    "Each card has a distinctive icon, bold title, and brief description.")

add_student_page("WhatsApp Image 2026-04-02 at 6.50.14 PM.jpeg", "Medical Excuses",
    "The Medical Excuses page allows students to manage and submit medical reports. The header shows 'Medical "
    "Excuses' with subtitle 'Manage and submit your medical reports' and a red '+ Submit New Excuse' button. "
    "A searchable list displays submitted excuses with title, date, and status badges: 'Severe Flu and Fever' "
    "(15 Oct 2025, Pending \u2014 amber), 'Dental Emergency' (05 Oct 2025, Approved \u2014 green), 'Food "
    "Poisoning' (01 Oct 2025, Rejected \u2014 red). The right panel shows 'No Excuse Selected' with a prompt "
    "to select an excuse to view its status and details, indicating a master-detail layout pattern.")

add_student_page("WhatsApp Image 2026-04-02 at 6.50.14 PM (1).jpeg", "Complaints",
    "The Complaints page enables students to 'Voice your concerns anonymously or officially'. The left panel "
    "shows a 'New Complaint' form with category selector chips (Academic \u2014 selected in red, Facilities, "
    "Financial, Administrative), a 'YOUR MESSAGE' textarea, and a red 'Submit Complaint' button. A "
    "confidentiality notice states complaints are addressed within 3-5 working days. The right 'Complaint "
    "History' panel shows past complaints: 'Library Noise Level' (RESOLVED, with official response) and "
    "'Portal Login Issue' (PENDING).")

add_student_page("WhatsApp Image 2026-04-02 at 6.50.14 PM (2).jpeg", "Academic Warnings",
    "The Academic Warnings page displays 'Important notifications regarding your academic standing'. Two "
    "warning cards are shown: an 'Attendance Warning' for CS402 (amber severity Medium, issued 12 Oct 2025) "
    "noting attendance dropped below 80%, and an 'Academic Warning' for Overall GPA (red severity High, "
    "issued 05 Sep 2025) noting GPA is 2.1 with probation risk. Each card has 'Request Appeal' and 'View "
    "Policy' buttons. An 'Academic Policy Note' at the bottom explains that warnings are issued automatically "
    "and can be contested within 7 days.")

add_student_page("WhatsApp Image 2026-04-02 at 6.50.14 PM (3).jpeg", "Official Requests",
    "The Official Requests page allows students to 'Order certificates and academic documents'. The 'Available "
    "Documents' section displays six document cards in a 3\u00d72 grid: Enrolment Certificate (100 EGP, 1-2 "
    "Days), Official Transcript (230 EGP, 3-5 Days), ID Card Replacement (150 EGP, Same Day), Course "
    "Description (50 EGP, 1-2 Days), Military Service Form (Free, 3 Days), and Graduation Statement (300 EGP, "
    "3 Days). The right 'Recent Orders' panel shows an 'Official Transcript' (COMPLETED, with Download E-Copy "
    "button) and 'Enrolment Certificate' (IN PROGRESS). A red 'Track Outside Requests?' widget allows tracking "
    "campus office requests by ID.")

doc.add_page_break()

# 2.5 Curriculum Management
doc.add_heading("2.5 Curriculum Management", level=2)

add_student_page("WhatsApp Image 2026-04-02 at 6.50.14 PM (4).jpeg", "Class Timetable",
    "The Class Timetable page shows 'Your weekly schedule for the current semester' with quick action buttons "
    "'Set Reminder' and 'Download PDF'. Four KPI cards show: This Week (11 sessions), Sections (8), Online (2), "
    "Tutorials (1). A color-coded weekly calendar grid displays courses across Monday-Friday from 8:00 to 16:00. "
    "Courses include CS431 Advanced Web Dev (blue), CS412 Database Systems (red), MATH301 Linear Algebra "
    "(orange), CS302 Machine Learning (purple), CS431L Web Dev Lab (teal), and ARBLEET Technical Writing (gray). "
    "A color legend at the bottom maps each course code to its assigned color.")

add_student_page("WhatsApp Image 2026-04-02 at 6.50.14 PM (5).jpeg", "Add / Drop Courses",
    "The Add/Drop Courses page lets students 'Manage your registered courses for the current semester'. A "
    "'Smart Recommendations' button in the top-right activates the AI recommendation system. Three KPI cards "
    "show: My Enrolled Courses (2), Total Credits (6), Available Courses (6). The left 'My Enrolled Courses' "
    "panel lists CS431 \u2014 Advanced Web Development (MWF 9:00-10:30, CS Lab 101) and MATH301 \u2014 Linear "
    "Algebra (TTH 8:00-9:30, Hall 201), each with a red 'Drop Course' button. The right 'Available Courses' "
    "panel shows searchable course cards with code, department, availability status, difficulty badges, "
    "instructor, schedule, credits, and enrollment capacity with progress bars.")

add_student_page("WhatsApp Image 2026-04-02 at 6.50.14 PM (6).jpeg", "Smart Recommendations Modal",
    "This modal overlay titled 'What\u2019s your goal?' appears when clicking 'Smart Recommendations'. It "
    "features a graduation cap icon and offers two goal-based course recommendation scenarios: 'Improve My GPA' "
    "\u2014 suggesting easier, high-impact courses to boost GPA (tagged 'High GPA Impact', 'Easier Workload'), "
    "and 'Complete My Requirements' \u2014 recommending courses with the most credit hours for faster graduation "
    "(tagged 'More Credits', 'Faster Graduation'). A note states 'You can always browse all courses after "
    "selecting a goal'. This AI-driven feature personalizes course selection based on student academic goals.")

add_student_page("WhatsApp Image 2026-04-02 at 6.50.14 PM (7).jpeg",
    "Add / Drop Courses \u2013 Improve GPA Goal",
    "This view shows the Add/Drop Courses page after selecting the 'Improve GPA' goal. The top banner displays "
    "'Goal: Improve GPA \u2014 Sorted by GPA impact \u2014 easiest courses first \u2014 3 courses recommended' "
    "with 'Show All' and 'Clear' options. An 'Improve GPA | Change' button in the top-right allows switching "
    "goals. Available courses now show a '\u2605 Recommended for you' blue banner and include a 'GPA+3.8' "
    "indicator showing predicted GPA impact. Courses are sorted by ease to maximize GPA improvement.")

add_student_page("WhatsApp Image 2026-04-02 at 6.50.15 PM.jpeg",
    "Add / Drop Courses \u2013 Complete Requirements Goal",
    "This view shows the Add/Drop page with the 'Complete Requirements' goal active. The banner shows "
    "'Goal: Complete Requirements \u2014 Sorted by credit hours \u2014 fastest path to graduation \u2014 3 "
    "courses recommended'. Available courses display green '\u2605 Recommended for you' banners and are sorted "
    "by credit hours to help the student maximize graduation progress. The 'Complete Requirements | Change' "
    "button allows switching to the alternative goal.")

doc.add_page_break()

# 2.6 My Grades
doc.add_heading("2.6 My Grades", level=2)

add_student_page("WhatsApp Image 2026-04-02 at 6.50.15 PM (1).jpeg", "My Grades",
    "The My Grades page provides a 'Detailed breakdown of your academic performance'. Four KPI cards display: "
    "Overall GPA (3.61), Avg Score (90.4%), Credits (16), and Courses (5). Below, each course is presented "
    "with its code, credit hours, title, instructor, overall percentage, letter grade, and trend indicator. "
    "CS431 Advanced Web Dev shows 95% (Grade A, Improving) with a 'GRADE BREAKDOWN' section visualizing "
    "individual components: Project 1 (90/100), Project 2 (86/100), Midterm Exam (82/100), Final Project "
    "(95/100) as horizontal progress bars. MATH301 Advanced Mathematics shows 88% (Grade B+, Improving) with "
    "similar component breakdowns. This detailed view gives students full transparency into their grades.")

doc.add_page_break()

# 2.7 Curriculum (My Curriculum)
doc.add_heading("2.7 My Curriculum", level=2)

add_student_page("WhatsApp Image 2026-04-02 at 6.50.15 PM (2).jpeg",
    "My Curriculum \u2013 Overview & First Year",
    "The My Curriculum page shows the student's complete degree roadmap for 'Bachelor of Science in Computer "
    "Science \u00b7 4 Years'. Five KPI indicators display: Overall Progress (60%), Completed (19 courses), "
    "In Progress (4 courses), Failed/Dropped (1), and Credits Done (52/87). A 'Curriculum Progress' bar "
    "visualizes 60% completion. A color-coded legend explains status indicators: Done (green), In Progress "
    "(blue), Upcoming (gray), Failed (red \u2717), Dropped (orange), and Has Prerequisites (link icon). "
    "The First Year section shows Fall 2022 and Spring 2023 semesters with courses like CS101, MATH101, "
    "ENG101 (all Done), and MATH102 (Failed \u2717) with prerequisite chains displayed.")

add_student_page("WhatsApp Image 2026-04-02 at 6.50.15 PM (3).jpeg",
    "My Curriculum \u2013 Second Year",
    "The curriculum view scrolled to show the Summer 2023 retake semester (MATH102 Calculus II, Failed \u2717, "
    "retaking from First Year) and Second Year courses. Fall 2023 shows CS201 Data Structures, CS203 OOP, "
    "MATH201 Linear Algebra, SOC201 Society & Tech (all Done). Spring 2024 shows CS211 Algorithms, CS214 OS "
    "Concepts, MATH204 Discrete Math, HUM206 Humanities (all Done). Each course card displays code, name, "
    "credits, prerequisite links, and completion status.")

add_student_page("WhatsApp Image 2026-04-02 at 6.50.15 PM (4).jpeg",
    "My Curriculum \u2013 Third Year",
    "The Third Year curriculum view shows Fall 2024 (CS301 Database Systems, CS302 Software Engineering, CS303 "
    "Networks, HUM301 Research Methods \u2014 all Done) and Spring 2025 (CS311 AI Fundamentals, CS321 Web "
    "Development, MATH301 Prob & Stats, ENG301 Technical Writing \u2014 all In Progress in blue). Summer "
    "periods show 'No courses \u2014 On Break'.")

add_student_page("WhatsApp Image 2026-04-02 at 6.50.15 PM (5).jpeg",
    "My Curriculum \u2013 Fourth Year & Failed Courses",
    "The Fourth Year shows upcoming courses: Fall 2025 (CS401 Graduation Project 6cr, CS411 Machine Learning, "
    "CS421 Cloud Computing) and Spring 2026 (CS431 Adv Web Dev, CS441 Security, CS451 Capstone) \u2014 all "
    "marked 'Upcoming' with prerequisite links. Below, a 'Failed & Dropped Courses Summary' table lists "
    "courses to retake, showing MATH102 Calculus II (3 credits, Failed \u2717, Retake In: Summer 2023). This "
    "comprehensive view helps students plan their remaining academic journey.")

doc.add_page_break()

# 2.8 Records
doc.add_heading("2.8 Records", level=2)

add_student_page("WhatsApp Image 2026-04-02 at 6.50.15 PM (6).jpeg", "Academic Records & Transcript",
    "The Records page shows 'Your official academic record and status'. The left profile card displays the "
    "student's avatar, name (Rawda Ayman), ID (21at41), Faculty (Engineering & Technology), Program (B.Sc. in "
    "Computer Science), Year (4th Year), Cumulative GPA (3.82/4.0), Credits (75/120), and Active Status "
    "(VALID). The right panel shows a 'Semester Transcript' table with columns: Semester, GPA, Credits, Honor. "
    "Five semesters are listed: Fall 2022 (3.8, 18cr, Dean's List), Spring 2023 (3.9, 18cr, Dean's List), "
    "Fall 2023 (3.7, 18cr), Spring 2024 (3.8, 18cr), Fall 2024 (3.9, 15cr, Dean's List). A 'Download PDF' "
    "button generates the official transcript document.")

doc.add_page_break()

# 2.9 AI Assistant
doc.add_heading("2.9 AI Assistant", level=2)

add_student_page("WhatsApp Image 2026-04-02 at 6.50.15 PM (7).jpeg", "AI Assistant \u2013 EduBot",
    "The AI Assistant page features 'EduBot', the EduSphere AI Academic Assistant. The interface has a red "
    "header banner displaying 'EduBot \u2022 AI Academic Assistant \u00b7 Online' with an 'EduSphere AI' badge. "
    "The chat area shows EduBot's welcome message: 'Hello! \U0001f44b I\u2019m EduBot, your AI academic "
    "assistant. How can I help you today?' (timestamped 08:20 PM). Quick action chips at the bottom provide "
    "common queries: 'My GPA', 'My Schedule', 'Attendance', 'Assignments', and 'Help'. A message input field "
    "with placeholder 'Type your message...' and a pink send button enable free-text interactions. This AI "
    "chatbot provides instant academic assistance, answering questions about grades, schedules, attendance, "
    "and general university information.")

doc.add_page_break()

# 2.10 Account Settings
doc.add_heading("2.10 Account Settings", level=2)

add_student_page("WhatsApp Image 2026-04-02 at 6.50.16 PM.jpeg", "Settings \u2013 Profile Tab",
    "The Account Settings page with the 'Profile' tab active. The left navigation shows six settings tabs: "
    "Profile (active, red), Security, Notifications, Appearance, Language, and Academic. The profile section "
    "displays the user's avatar, name (Rawda Ayman), and email (rawda.ayman@edusphere.edu). A notice states "
    "'Your profile information is managed by the university and cannot be edited.' Read-only fields show: "
    "Full Name, Student ID (21at41), Email, Phone (+20 100 000 0000), Faculty (Engineering & Technology), "
    "and Program (B.Sc. Computer Science).")

add_student_page("WhatsApp Image 2026-04-02 at 6.50.16 PM (1).jpeg", "Settings \u2013 Security Tab",
    "The Security tab shows a 'Change Password' section with fields for Current Password, New Password, and "
    "Confirm Password (all masked), plus a red 'Save Password' button. Below, an 'Active Sessions' section "
    "displays devices currently signed in: 'Chrome on Windows' (IP: 192.168.1.5, labeled 'THIS DEVICE') and "
    "'Mobile - Safari iOS' (IP: 41.234.12.80, 2 hours ago) with a red 'Revoke' button to terminate the remote "
    "session.")

add_student_page("WhatsApp Image 2026-04-02 at 6.50.16 PM (2).jpeg",
    "Settings \u2013 Security Tab (Two-Factor Authentication)",
    "This scrolled view of the Security tab shows the 'Two-Factor Authentication' section for adding extra "
    "security. Two toggle switches are available: 'Authenticator App' (Use Google Authenticator or similar "
    "\u2014 enabled) and 'SMS Verification' (Receive codes via text message \u2014 enabled). Both toggles "
    "are shown in the active (red) state, indicating the student has enabled both 2FA methods for maximum "
    "account protection.")

add_student_page("WhatsApp Image 2026-04-02 at 6.50.16 PM (3).jpeg", "Settings \u2013 Notifications Tab",
    "The Notifications tab shows 'Notification Preferences \u2014 Choose which notifications you want to "
    "receive.' Five toggle switches control: Email Notifications (Receive updates via email \u2014 enabled), "
    "Push Notifications (Browser alerts \u2014 enabled), Grade Updates (When new grades are posted \u2014 "
    "enabled), Assignment Reminders (Before due dates \u2014 enabled), and Attendance Alerts (Absence warnings "
    "\u2014 enabled). All toggles are shown in the active (red) state.")

add_student_page("WhatsApp Image 2026-04-02 at 6.50.16 PM (4).jpeg", "Settings \u2013 Appearance Tab",
    "The Appearance tab under 'THEME' offers two theme options: 'Dark Mode' (moon icon) and 'Light Mode' "
    "(sun icon, marked '\u2713 Active'). The clean card-based selection allows students to switch between "
    "light and dark visual themes to suit their preference and reduce eye strain during extended sessions.")

add_student_page("WhatsApp Image 2026-04-02 at 6.50.16 PM (5).jpeg", "Settings \u2013 Language Tab",
    "The Language tab displays five language options in a card grid layout: English (GB EN, selected with a "
    "red checkmark), Arabic (EG AR), Russian (RU RU), French (FR FR), and German (DE DE). Each card shows the "
    "country code and language abbreviation, providing multilingual support for the international student body.")

add_student_page("WhatsApp Image 2026-04-02 at 6.50.16 PM (6).jpeg", "Settings \u2013 Academic Tab",
    "The Academic tab displays key academic information in a 3\u00d72 card grid: Current GPA (3.82/4.0 in "
    "green), Credits Earned (75/120 in blue), Academic Standing (Good Standing in green), Expected Graduation "
    "(June 2026), Major (Computer Science), and Minor (Mathematics). This provides a quick academic snapshot "
    "accessible from the settings page.")

doc.add_page_break()

# 2.11 Get Support & Sign Out
doc.add_heading("2.11 Get Support & Sign Out", level=2)

add_student_page("WhatsApp Image 2026-04-02 at 6.50.17 PM.jpeg", "Get Support Modal",
    "The Get Support modal, triggered by clicking the 'Get Support' button in the sidebar's 'Need Help?' "
    "section. It displays 'EduSphere IT & Academic Support' and provides a support ticket form with an "
    "'ISSUE TYPE' selector offering 'Technical', 'Academic', and 'Other' category chips. A 'DESCRIBE YOUR "
    "ISSUE' textarea (placeholder 'Describe your problem...') allows detailed issue description. The red "
    "'Submit Ticket' button sends the support request to the IT team. A close (\u00d7) button dismisses the "
    "modal.")

add_student_page("WhatsApp Image 2026-04-02 at 6.50.17 PM (1).jpeg", "Sign Out Confirmation",
    "The Sign Out confirmation dialog appears when clicking 'Sign Out' in the sidebar. It features a door/exit "
    "icon and asks 'Sign Out?' with the message 'You\u2019ll be signed out of EduSphere. Make sure you\u2019ve "
    "saved your work.' Two buttons are provided: 'Cancel' (outlined) to stay signed in, and 'Sign Out' (red) "
    "to confirm logout. This confirmation step prevents accidental session termination.")

doc.add_page_break()

# 2.12 Dark Mode
doc.add_heading("2.12 Dark Mode", level=2)

add_student_page("WhatsApp Image 2026-04-02 at 6.50.17 PM (3).jpeg", "Student Dashboard \u2013 Dark Mode",
    "This screenshot shows the Student Dashboard in Dark Mode, demonstrating the platform's full dark theme "
    "implementation. The entire interface \u2014 sidebar, top bar, KPI cards, class schedule, and grades panel "
    "\u2014 uses a dark color palette with light text for readability. The dark mode toggle (sun icon) in the "
    "top navigation bar switches between themes. All UI elements maintain their functionality and visual "
    "hierarchy while using darker backgrounds and adjusted color contrast, providing a comfortable viewing "
    "experience in low-light environments and reducing eye strain during extended study sessions.")

doc.add_page_break()

# ═══════════════════════════════════════════════
#  3. ADMINISTRATOR ROLE
# ═══════════════════════════════════════════════
doc.add_heading("3. Administrator Role", level=1)
doc.add_paragraph(
    "The Administrator has full access to all system modules. This includes managing students, "
    "faculty members, courses, registrations, scheduling, academic performance tracking, "
    "reports generation, and system administration. The admin dashboard provides a comprehensive "
    "overview of the entire institution's operations."
)

doc.add_heading("3.1 Admin Dashboard", level=2)
add_page("dashboard_overview_*.png", "Admin Dashboard \u2013 Overview",
    "This is the main Admin Dashboard, titled 'University Overview'. It presents a comprehensive "
    "summary of the institution's current state through four KPI cards: Total Students, Courses, "
    "Faculty Members, and Pending Requests. Each card displays real-time data fetched from the backend. "
    "The dashboard includes a semester selector, 'Export Report' and 'Run Conflict Scan' buttons. "
    "Interactive charts visualize student distribution by level (bar chart) and program (donut chart), "
    "Request Status Overview, a 'Recent Requests' feed, and a 'Platform Overview' widget.")
doc.add_page_break()

doc.add_heading("3.2 Student Management", level=2)
add_page("students_list_page_*.png", "Students List Page",
    "The Students Management page with a searchable data table (Student ID, Name, Program, Level, "
    "Credit Hours), search bar, Filters, Export buttons, and action icons (view, edit, delete). "
    "The '+ Add Student' button opens the creation modal.")
add_page("add_student_modal_*.png", "Add New Student Modal",
    "Modal for registering a new student with User dropdown, Program dropdown, Level (default 1), "
    "and Credit Hours (default 0) fields. Includes Cancel and Add Student buttons.")
add_page("v2_edit_faculty_modal_*.png", "Edit Student Modal",
    "Edit modal pre-populated with the student's current program, level, and credit hours for modification.")
add_page("v2_delete_faculty_dialog_*.png", "Delete Student Confirmation",
    "Confirmation dialog warning about permanent deletion with Cancel and Delete buttons.")
add_page("student_affairs_detail_*.png", "Student Profile Page",
    "Detailed student profile with personal details, academic information, enrollment status, and performance history.")
doc.add_page_break()

doc.add_heading("3.3 Faculty Management", level=2)
add_page("v2_faculty_list_*.png", "Faculty Members List",
    "Faculty management page with card-based layout, four summary cards (Total, Full-Time, Part-Time, "
    "Departments), search/filter/export options, and faculty cards with initials, name, ID, employment badge, "
    "department, and email. Hover reveals action buttons.")
add_page("v2_add_faculty_modal_*.png", "Add Faculty Member Modal",
    "Registration modal with User dropdown, Department dropdown, and Employment Type (Full-Time/Part-Time) selector.")
add_page("v2_edit_faculty_modal_*.png", "Edit Faculty Member Modal",
    "Edit modal for modifying department and employment type with pre-filled values.")
add_page("v2_delete_faculty_dialog_*.png", "Delete Faculty Member Confirmation",
    "Deletion confirmation dialog with warning message, Cancel and Delete buttons.")
add_page("v2_faculty_detail_*.png", "Faculty Member Detail Page",
    "Detailed faculty profile with personal details, credentials, department, employment type, schedule, and courses.")
doc.add_page_break()

doc.add_heading("3.4 Course Management", level=2)
add_page("v2_semesters_page_*.png", "Semesters Management",
    "Semester configuration page with name, dates, and active/inactive status toggle.")
add_page("v2_courses_page_*.png", "Course Catalog",
    "Master course registry with code, title, credit hours, department columns and CRUD operations.")
add_page("v2_sections_page_*.png", "Sections Management",
    "Section management with instructor assignment, capacity, time slots, and room assignments.")
add_page("v2_prerequisites_page_*.png", "Prerequisites Configuration",
    "Course prerequisite and dependency chain configuration interface.")
doc.add_page_break()

doc.add_heading("3.5 Registration & Operations", level=2)
add_page("registration_hub_*.png", "Registration Hub",
    "Central enrollment control panel with registration period overview, dates, limits, and conflict resolution.")
add_page("registration_requests_*.png", "Registration Requests",
    "Student registration requests list with approval/rejection workflow and batch processing.")
doc.add_page_break()

doc.add_heading("3.6 Student Affairs Module", level=2)
add_page("student_affairs_list_*.png", "Student Affairs Staff List",
    "Student affairs personnel directory with assignments and workload distribution.")
add_page("student_requests_*.png", "Student Requests Management",
    "Centralized non-registration request management with review, approve/reject workflow.")
add_page("student_phones_*.png", "Student Contact Directory",
    "Student phone directory with search and filtering for communication needs.")
doc.add_page_break()

doc.add_heading("3.7 Scheduling", level=2)
add_page("class_schedules_*.png", "Class Schedules",
    "Weekly timetable view for conflict detection and time slot optimization.")
add_page("room_management_*.png", "Room Management",
    "Campus facility administration with room types, capacities, and availability.")
add_page("academic_calendar_*.png", "Academic Calendar",
    "Official academic schedule with semester dates, exams, deadlines, and holidays.")
doc.add_page_break()

doc.add_heading("3.8 Academic Performance", level=2)
add_page("grades_management_*.png", "Grades Management",
    "Institutional grading overview with submission tracking and distribution patterns.")
add_page("gpa_overview_*.png", "GPA Overview",
    "GPA analytics dashboard with trends by program, department, and level.")
add_page("transcripts_*.png", "Transcripts Management",
    "Transcript request processing and official document generation.")
doc.add_page_break()

doc.add_heading("3.9 Reports & Analytics", level=2)
add_page("enrollment_report_page_1775058*.png", "Enrollment Report",
    "Enrollment analytics with program, level, gender breakdowns and trend analysis.")
add_page("performance_report_page_1775058*.png", "Academic Performance Report",
    "Performance analytics with department comparisons, pass rates, and GPA distributions.")
doc.add_page_break()

doc.add_heading("3.10 System Administration", level=2)
add_page("user_management_page_*.png", "User Management",
    "User account management with ID, name, email, verification status, and CRUD operations.")
add_page("add_user_modal_*.png", "Add New User Modal",
    "User creation modal with name, email, password, and role assignment fields.")
add_page("role_management_page_*.png", "Role Management",
    "Role and permission management with role cards showing assigned user counts.")
add_page("system_settings_page_*.png", "System Settings",
    "Platform configuration with General, Appearance, Security, and Notifications tabs.")
doc.add_page_break()

# ═══════════════════════════════════════════════
#  4. INSTRUCTOR ROLE
# ═══════════════════════════════════════════════
doc.add_heading("4. Instructor Role", level=1)
doc.add_paragraph(
    "Instructors have access to features related to their teaching responsibilities including "
    "course management, student performance tracking, grade submission, and institutional reports."
)

doc.add_heading("4.1 Instructor Dashboard", level=2)
add_page("instructor_dashboard_1775058*.png", "Instructor Dashboard",
    "Personalized dashboard with active courses, student capacity, teaching days, session duration metrics, "
    "'My Courses' table, 'My Schedule' panel, 'Sessions by Type' chart, and grade export tools.")

doc.add_heading("4.2 Instructor Profile", level=2)
add_page("instructor_profile_1775058*.png", "Instructor Profile",
    "Faculty profile page with name, title, email, department, and employment type information.")

doc.add_heading("4.3 My Courses", level=2)
add_page("instructor_courses_1775059*.png", "My Courses Page",
    "Course listing with code, title, section, schedule, enrolled count, capacity, and workspace navigation.")
add_page("instructor_course_details_*.png", "Course Workspace",
    "Detailed course management with Students & Grades, Components, Materials, and Attendance tabs.")

doc.add_heading("4.4 Grade Management", level=2)
add_page("instructor_grades_1775059*.png", "Instructor Grades View",
    "Grade entry interface with student lists, grade validation, and submission progress tracking.")

doc.add_heading("4.5 Instructor Reports", level=2)
add_page("instructor_enrollment_report_*.png", "Instructor Enrollment Report",
    "Department-level enrollment insights with trends and section fill rates.")
add_page("instructor_performance_report_*.png", "Instructor Performance Report",
    "Course performance analytics with pass rates, grade distributions, and GPA trends.")
doc.add_page_break()

# ═══════════════════════════════════════════════
#  5. STUDENT AFFAIRS STAFF ROLE
# ═══════════════════════════════════════════════
doc.add_heading("5. Student Affairs Staff Role", level=1)
doc.add_paragraph(
    "Student Affairs staff focus on student support through request processing, registration "
    "management, and maintaining accurate student records."
)

doc.add_heading("5.1 Student Affairs Dashboard", level=2)
add_page("student_affairs_dashboard_*.png", "Student Affairs Dashboard",
    "Operations dashboard with Pending/Approved/Active/Rejected metrics, Request Status Overview, "
    "Requests by Type chart, Recent Requests table, and Quick Actions panel.")

doc.add_heading("5.2 Student Management (Affairs View)", level=2)
add_page("student_affairs_students_list_*.png", "Students List (Affairs View)",
    "Student directory with search/filter capabilities and role-limited actions.")

doc.add_heading("5.3 Student Affairs Operations", level=2)
add_page("student_affairs_profile_info_*.png", "Student Affairs Staff Directory",
    "Department staff directory with roles, contact info, and assignment areas.")
add_page("student_affairs_student_requests_*.png", "Student Requests Processing",
    "Request processing interface with status updates and workflow management.")
add_page("student_affairs_phone_directory_*.png", "Student Phone Directory (Affairs View)",
    "Contact lookup for student communication and administrative follow-ups.")

doc.add_heading("5.4 Registration Management (Affairs View)", level=2)
add_page("student_affairs_registration_management_loaded_*.png", "Registration Management",
    "Enrollment assistance with manual overrides and registration progress monitoring.")
add_page("student_affairs_requests_general_*.png", "General Requests Page",
    "Comprehensive request view with date, type, and status filtering.")

doc.add_heading("5.5 Student Affairs Reports", level=2)
add_page("student_affairs_enrollment_reports_*.png", "Enrollment Report (Affairs View)",
    "Enrollment analytics for student services planning and resource allocation.")
add_page("student_affairs_performance_reports_*.png", "Performance Report (Affairs View)",
    "Academic performance insights for advising and early intervention programs.")

# ═══════════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════════
doc.save(output_path)
print(f"\n\u2705 Document saved successfully to:\n   {output_path}")
print(f"   Total figures: {fig_counter[0] - 1}")
